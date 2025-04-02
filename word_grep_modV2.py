import os
import re
from docx import Document

output_file = "output.txt"

def word_grep(filename, path, search_term, output_file, write_line_text=True, whole_word=False, case_sensitive=False):
    file_path = os.path.join(path, filename)

    if filename.lower().endswith('.docx'):
        try:
            doc = Document(file_path)
        except Exception as e:
            print(f"Error reading Word file {filename}: {e}")
            return 0
        results = search_word_document(doc, search_term, whole_word, case_sensitive)
    else:
        # Handle other types of files (Excel, PowerPoint, etc.) if needed
        results = []

    # Write results to the output file
    with open(output_file, 'a', encoding='utf-8') as output:
        if results:
            filename_line = fr"find file: {path}\{filename}  -----"
            output.write(f"{filename_line}\n")

            # Write results line by line
            if write_line_text:
                for write_line in results:
                    try:
                        output.write(f"{write_line}\n")
                    except Exception as e:
                        print(f"Error writing line to output file: {e}")
                output.write("\n")

            return 1
        else:
            return 0

def search_word_document(doc, search_term, whole_word, case_sensitive):
    results = []

    # Regular expression options
    flags = 0
    if not case_sensitive:
        flags |= re.IGNORECASE

    # Search normal paragraphs
    for paragraph_index, paragraph in enumerate(doc.paragraphs):
        if re.search(re.escape(search_term) if whole_word else search_term, paragraph.text, flags=flags):
            result_info = f'Paragraph: {paragraph_index + 1}, Text: {paragraph.text}'
            results.append(result_info)

    # Search tables
    for table_index, table in enumerate(doc.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                if re.search(re.escape(search_term) if whole_word else search_term, cell.text, flags=flags):
                    result_info = f'Table: {table_index + 1}, Row: {row_index + 1}, Cell: {cell_index + 1}, Text: {cell.text}'
                    results.append(result_info)

    return results

if __name__ == "__main__":
    word_grep(r"wd-spectools-word-sample-04.docx", r"C:\Users\xzyoi\Downloads" ,"Status" , output_file)
